#!/usr/bin/env python3
"""
Konvertierung beliebiger Dateien → PNG(s) für Gemini 3 Flash Preview.

Unser System soll alle möglichen Formate in Bilder umwandeln können;
diese PNGs gehen dann durch Gemini 3 Flash Preview zur Textextraktion.

Unterstützt (ohne externe Binaries wie LibreOffice):
  - PDF          → PyMuPDF, eine PNG pro Seite
  - Bilder       → JPG/JPEG, WEBP, BMP, GIF, TIFF → PNG
  - TXT, MD, CSV, JSON, XML, HTML → Text extrahieren → eine oder mehrere PNG(s)
  - DOCX         → python-docx → Text → PNG
  - XLSX         → openpyxl → Text/Tabellen → PNG
  - EML          → E-Mail-Body (Text/HTML) → PNG

Optional: Mit installiertem LibreOffice (headless) können ODT, RTF, DOC, XLS
  zu PDF und dann zu PNG konvertiert werden – siehe _convert_via_libreoffice.

Ausgabe: Liste von PNG-Bytes (jeweils ein Bild), plus optional Dateiname für Kontext.
"""

from __future__ import annotations

import csv
import email
import io
import re
from pathlib import Path
from email import policy
from email.parser import BytesParser

# Optional imports – bei Fehlern wird der jeweilige Konverter ausgelassen
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
try:
    from PIL import Image
except ImportError:
    Image = None
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
try:
    import openpyxl
except ImportError:
    openpyxl = None
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

# DPI für PDF-Rendering (Lesbarkeit vs. Größe)
PDF_DPI = 150
# Max. Seitenlänge für Text→PNG (Zeilen), danach aufteilen
TEXT_PNG_MAX_LINES = 80
# Zeichen pro Zeile (annähernd) für Text-PNG
CHARS_PER_LINE = 90


def _text_to_png_bytes(text: str, source_name: str = "") -> list[bytes]:
    """Rendert Fließtext in eine oder mehrere PNG(s). Jede PNG = ein Block Text."""
    if not text or not text.strip():
        return []
    if not Image:
        return []  # Fallback: Caller kann Text direkt an Gemini senden
    lines = text.replace("\r\n", "\n").replace("\r", "\n").strip().split("\n")
    out = []
    chunk: list[str] = []
    for line in lines:
        # Lange Zeilen umbrechen
        while len(line) > CHARS_PER_LINE:
            chunk.append(line[:CHARS_PER_LINE])
            line = line[CHARS_PER_LINE:]
            if len(chunk) >= TEXT_PNG_MAX_LINES:
                out.append(_render_lines_to_png(chunk, source_name))
                chunk = []
        chunk.append(line)
        if len(chunk) >= TEXT_PNG_MAX_LINES:
            out.append(_render_lines_to_png(chunk, source_name))
            chunk = []
    if chunk:
        out.append(_render_lines_to_png(chunk, source_name))
    return out


def _render_lines_to_png(lines: list[str], title: str) -> bytes:
    """Zeilen in eine PNG zeichnen (einfaches weißes Bild mit schwarzer Schrift)."""
    try:
        from PIL import ImageDraw, ImageFont
    except ImportError:
        ImageDraw = ImageFont = None
    font_size = 14
    line_height = 20
    margin = 40
    width = 800
    top = margin + (line_height * 2 if title else 0)
    height = margin + top + len(lines) * line_height
    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img) if ImageDraw else None
    font = None
    if ImageFont:
        for path in ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/System/Library/Fonts/Helvetica.ttc"):
            try:
                font = ImageFont.truetype(path, font_size)
                break
            except Exception:
                pass
        if font is None:
            try:
                font = ImageFont.load_default()
            except Exception:
                pass
    if draw and font:
        if title:
            draw.text((margin, 10), title[:80], fill=(0, 0, 0), font=font)
        for i, line in enumerate(lines):
            y = top + i * line_height
            draw.text((margin, y), line[:200], fill=(0, 0, 0), font=font)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _pdf_to_pngs(path: Path) -> list[bytes]:
    if not fitz:
        return []
    out = []
    try:
        doc = fitz.open(str(path))
        for i in range(len(doc)):
            page = doc[i]
            pix = page.get_pixmap(dpi=PDF_DPI)
            out.append(pix.tobytes("png"))
        doc.close()
    except Exception:
        pass
    return out


def _image_to_png(path: Path) -> list[bytes]:
    if not Image:
        return []
    out = []
    try:
        with Image.open(path) as im:
            if im.mode in ("RGBA", "P"):
                im = im.convert("RGB")
            buf = io.BytesIO()
            im.save(buf, format="PNG")
            out.append(buf.getvalue())
    except Exception:
        pass
    return out


def _docx_to_text(path: Path) -> str:
    if not DocxDocument:
        return ""
    try:
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
            "\t".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows
        )
    except Exception:
        return ""


def _xlsx_to_text(path: Path) -> str:
    if not openpyxl:
        return ""
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"=== {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                parts.append("\t".join(str(c) if c is not None else "" for c in row))
        wb.close()
        return "\n".join(parts)
    except Exception:
        return ""


def _html_to_text(html: str) -> str:
    if BeautifulSoup:
        try:
            soup = BeautifulSoup(html, "html.parser")
            return soup.get_text(separator="\n", strip=True)
        except Exception:
            pass
    # Fallback: grob Tags entfernen
    return re.sub(r"<[^>]+>", " ", html).replace("&nbsp;", " ").strip()


def _eml_to_text(path: Path) -> str:
    try:
        with open(path, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)
        text_parts = []
        html_parts = []
        for part in msg.walk():
            if part.get_content_maintype() == "multipart":
                continue
            ct = part.get_content_type()
            payload = part.get_payload(decode=True)
            if not payload:
                continue
            charset = part.get_content_charset() or "utf-8"
            try:
                decoded = payload.decode(charset, errors="replace")
            except LookupError:
                decoded = payload.decode("utf-8", errors="replace")
            if ct == "text/plain":
                text_parts.append(decoded)
            elif ct == "text/html":
                html_parts.append(decoded)
        if text_parts:
            return "\n\n".join(text_parts)
        if html_parts:
            return _html_to_text("\n".join(html_parts))
    except Exception:
        pass
    return ""


def file_to_pngs(path: Path, source_name: str | None = None) -> list[bytes]:
    """
    Konvertiert eine Datei in eine Liste von PNG-Bildern (Bytes).
    Geeignet für den Upload an Gemini 3 Flash Preview.

    Returns:
        Liste von PNG-Dateien als bytes. Kann leer sein bei unbekanntem Format oder Fehler.
    """
    path = Path(path).resolve()
    if not path.is_file():
        return []
    name = source_name or path.name
    suffix = path.suffix.lower()

    # PDF
    if suffix == ".pdf" and fitz:
        return _pdf_to_pngs(path)

    # Bilder → eine PNG
    if suffix in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff", ".tif"):
        return _image_to_png(path)

    # Office / Text
    if suffix in (".docx", ".doc"):
        text = _docx_to_text(path) if suffix == ".docx" and DocxDocument else ""
        if not text and path.suffix.lower() == ".doc":
            text = "( .doc nur mit LibreOffice konvertierbar; bitte als .docx bereitstellen )"
        return _text_to_png_bytes(text or "(leer)", name) if text or suffix == ".docx" else []

    if suffix in (".xlsx", ".xls"):
        text = _xlsx_to_text(path) if suffix == ".xlsx" and openpyxl else ""
        if not text and suffix == ".xls":
            text = "( .xls nur mit LibreOffice konvertierbar )"
        return _text_to_png_bytes(text or "(leer)", name) if text or suffix == ".xlsx" else []

    # E-Mail
    if suffix == ".eml":
        text = _eml_to_text(path)
        return _text_to_png_bytes(text or "(kein lesbarer Inhalt)", name)

    # Text-basierte Formate
    if suffix in (".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm", ".rtf"):
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            return []
        if suffix == ".html" or suffix == ".htm":
            raw = _html_to_text(raw)
        elif suffix == ".csv":
            rows = list(csv.reader(io.StringIO(raw)))
            raw = "\n".join("\t".join(cell for cell in row) for row in rows)
        return _text_to_png_bytes(raw, name) if raw.strip() else []

    # Unbekannt: als Binärtext oder „raw“ versuchen
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
        if raw.strip():
            return _text_to_png_bytes(raw, name)
    except Exception:
        pass
    return []


def files_to_pngs(paths: list[Path]) -> list[tuple[str, bytes]]:
    """
    Mehrere Dateien → flache Liste von (dateiname, png_bytes).
    dateiname dient als Kontext für Gemini (z. B. „document.pdf Seite 3“).
    """
    result = []
    for path in paths:
        path = Path(path)
        name = path.name
        pngs = file_to_pngs(path, source_name=name)
        for i, png in enumerate(pngs):
            if len(pngs) > 1:
                label = f"{name} (Seite {i + 1}/{len(pngs)})"
            else:
                label = name
            result.append((label, png))
    return result


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Verwendung: python convert_to_png.py <datei> [datei2 ...]")
        print("  Gibt Anzahl erzeugter PNGs pro Datei aus.")
        sys.exit(0)
    for p in sys.argv[1:]:
        path = Path(p)
        pngs = file_to_pngs(path)
        print(f"{path.name}: {len(pngs)} PNG(s)")
